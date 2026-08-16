import re
import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tcPr = cell._tc.get_or_add_tcPr()
    borders_elm = parse_xml(f'<w:tcBorders {nsdecls("w")}/>')
    
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for side, border in borders.items():
        if border:
            val, sz, space, color = border
            b_elm = parse_xml(f'<w:{side} {nsdecls("w")} w:val="{val}" w:sz="{sz}" w:space="{space}" w:color="{color}"/>')
            borders_elm.append(b_elm)
        else:
            b_elm = parse_xml(f'<w:{side} {nsdecls("w")} w:val="none"/>')
            borders_elm.append(b_elm)
    tcPr.append(borders_elm)

def make_table_row_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = parse_xml(f'<w:tblHeader {nsdecls("w")}/>')
    cantSplit = parse_xml(f'<w:cantSplit {nsdecls("w")}/>')
    trPr.append(tblHeader)
    trPr.append(cantSplit)

def make_table_row_cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    cantSplit = parse_xml(f'<w:cantSplit {nsdecls("w")}/>')
    trPr.append(cantSplit)

def format_runs(p, text, default_color=None, default_bold=False, default_italic=False, default_size=None):
    # Regex to handle bold-italic, bold, italic, code, markdown links, and appendix pills
    tokens = re.split(r'(`\[See Appendix [^\]]+\]`|\[See Appendix [^\]]+\]|\*\*.*?\*\*|\*.*?\*|`.*?`|\[.*?\]\(.*?\))', text)
    for token in tokens:
        if not token:
            continue
        
        # Check appendix pill
        m_app = re.match(r'`?\[See Appendix ([^\]]+)\]`?', token)
        if m_app:
            run = p.add_run(f" [See Appendix {m_app.group(1)}] ")
            run.font.size = Pt(8.5) if default_size is None else Pt(default_size - 0.5)
            run.font.bold = True
            run.font.color.rgb = RGBColor(26, 115, 232) # #1A73E8
            continue
        
        # Check markdown link
        m_link = re.match(r'\[(.*?)\]\((.*?)\)', token)
        if m_link:
            link_text = m_link.group(1)
            run = p.add_run(link_text)
            if default_size:
                run.font.size = Pt(default_size)
            run.font.color.rgb = RGBColor(23, 78, 166) # Deep link blue
            run.font.underline = True
            continue
        
        # Check bold
        if token.startswith('**') and token.endswith('**') and len(token) >= 4:
            run = p.add_run(token[2:-2])
            run.font.bold = True
            if default_italic:
                run.font.italic = True
            if default_size:
                run.font.size = Pt(default_size)
            if default_color:
                run.font.color.rgb = default_color
            continue
            
        # Check italic
        if token.startswith('*') and token.endswith('*') and len(token) >= 2:
            run = p.add_run(token[1:-1])
            run.font.italic = True
            if default_bold:
                run.font.bold = True
            if default_size:
                run.font.size = Pt(default_size)
            if default_color:
                run.font.color.rgb = default_color
            continue

        # Check code
        if token.startswith('`') and token.endswith('`') and len(token) >= 2:
            run = p.add_run(token[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9) if default_size is None else Pt(default_size - 1)
            run.font.color.rgb = RGBColor(184, 18, 55)
            continue
            
        # Normal text
        run = p.add_run(token)
        if default_bold:
            run.font.bold = True
        if default_italic:
            run.font.italic = True
        if default_size:
            run.font.size = Pt(default_size)
        if default_color:
            run.font.color.rgb = default_color

def create_callout_box(doc, title, text_items, fill_hex="F8F9FA", border_color="1A73E8"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    tbl.columns[0].width = Inches(6.9)
    
    cell = tbl.cell(0, 0)
    set_cell_background(cell, fill_hex)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=160)
    set_cell_borders(cell, left=('single', '24', '0', border_color),
                           top=('single', '4', '0', 'DADCE0'),
                           bottom=('single', '4', '0', 'DADCE0'),
                           right=('single', '4', '0', 'DADCE0'))
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    r_title = p.add_run(title)
    r_title.font.bold = True
    r_title.font.size = Pt(10.5)
    r_title.font.color.rgb = RGBColor(26, 115, 232)
    
    for item in text_items:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.line_spacing = 1.15
        format_runs(p2, item, default_size=9.5)
        
    p_end = doc.add_paragraph()
    p_end.paragraph_format.space_before = Pt(0)
    p_end.paragraph_format.space_after = Pt(4)

def calculate_column_weights(headers):
    # Determine custom proportional weights based on table column count and header semantics
    cols_count = len(headers)
    hdr_lower = [h.lower() for h in headers]
    
    # 1. 33-Account Matrix (6 cols)
    if 'account name' in hdr_lower[0] and 'entry wedge' in hdr_lower[4]:
        return [1.3, 0.8, 0.9, 0.8, 1.8, 1.8]
    
    # 2. Hyperscaler Comparison (5 cols)
    if 'evaluation dimension' in hdr_lower[0] and 'hyperscaler winner' in hdr_lower[4]:
        return [1.2, 1.6, 1.3, 1.3, 1.8]
        
    # 3. HPC Benchmarks (5 cols)
    if 'performance metric' in hdr_lower[0]:
        return [1.4, 1.6, 1.4, 1.4, 1.6]
        
    # 4. 6-Agent Architecture (5 cols)
    if 'reference agent pattern' in hdr_lower[0]:
        return [1.4, 1.4, 1.6, 1.4, 1.6]
        
    # 5. ISV Foundation Model (5 cols)
    if 'isv partner name' in hdr_lower[0]:
        return [1.2, 1.6, 1.2, 1.4, 1.6]
        
    # 6. Revenue Trajectory (7 cols: Year, AI/ML, HPC, Data, Sovereign, CCUS, Total)
    if 'year' in hdr_lower[0] and cols_count == 7:
        return [0.7, 1.3, 1.1, 1.1, 1.2, 0.8, 0.8]
        
    # 7. Appendix A.2 Workload Breakdown (8 cols)
    if 'workload category' in hdr_lower[0] and cols_count == 8:
        return [1.8, 0.6, 0.6, 0.6, 0.6, 0.6, 0.8, 1.8]
        
    # 8. Digital Maturity Matrix (6 cols)
    if 'operator category' in hdr_lower[0] and cols_count == 6:
        return [1.1, 0.8, 1.4, 1.1, 1.3, 1.7]
        
    # 9. Startup Accelerator (4 cols)
    if 'startup name' in hdr_lower[0] and cols_count == 4:
        return [1.4, 1.8, 1.6, 2.2]
        
    # 10. Sovereign Topology Layer (4 cols)
    if 'topology layer' in hdr_lower[0] and cols_count == 4:
        return [1.5, 2.0, 1.8, 1.9]
        
    # 11. Project Interchange (4 cols)
    if 'dimension' in hdr_lower[0] and cols_count == 4:
        return [1.3, 2.0, 2.0, 1.9]
        
    # 12. DeepMind Energy Lab (4 cols)
    if 'engagement tier' in hdr_lower[0] and cols_count == 4:
        return [1.4, 1.5, 2.3, 2.0]
        
    # 13. CCUS Consortium (7 cols)
    if 'consortium target' in hdr_lower[0] and cols_count == 7:
        return [1.2, 1.2, 1.1, 0.7, 0.9, 0.9, 2.0]
        
    # 14. E.1 90-Day Metrics (6 cols)
    if 'field enablement metric' in hdr_lower[0] and cols_count == 6:
        return [2.0, 0.6, 0.6, 0.6, 0.6, 1.8]
        
    # 15. E.2 Kit Registry (4 cols)
    if 'kit name' in hdr_lower[0] and cols_count == 4:
        return [1.4, 1.3, 2.7, 1.2]
        
    # 16. Appendix F References (7 cols)
    if 'ref id' in hdr_lower[0] and cols_count == 7:
        return [0.7, 1.5, 1.1, 0.8, 0.7, 1.8, 1.0]

    # Default equal weights
    return [1.0] * cols_count

def render_markdown_table(doc, table_rows):
    if len(table_rows) < 2:
        return
        
    header_cells = [c.strip() for c in table_rows[0].strip('|').split('|')]
    data_rows = []
    for r in table_rows[2:]:
        cells = [c.strip() for c in r.strip('|').split('|')]
        while len(cells) < len(header_cells):
            cells.append("")
        data_rows.append(cells[:len(header_cells)])
        
    cols_count = len(header_cells)
    rows_count = len(data_rows) + 1
    
    tbl = doc.add_table(rows=rows_count, cols=cols_count)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    total_avail_width = 6.9 # inches
    weights = calculate_column_weights(header_cells)
    sum_weights = sum(weights)
    col_widths = [Inches(total_avail_width * (w / sum_weights)) for w in weights]
    
    # Header Row
    hdr_row = tbl.rows[0]
    make_table_row_header(hdr_row)
    for j, h_text in enumerate(header_cells):
        cell = hdr_row.cells[j]
        cell.width = col_widths[j]
        set_cell_background(cell, "1A73E8") # Google Blue
        set_cell_margins(cell, top=90, bottom=90, left=100, right=100)
        set_cell_borders(cell, top=('single', '6', '0', '1A73E8'),
                               bottom=('single', '12', '0', '174EA6'),
                               left=('single', '4', '0', 'DADCE0'),
                               right=('single', '4', '0', 'DADCE0'))
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        format_runs(p, h_text, default_bold=True, default_color=RGBColor(255, 255, 255), default_size=8.5)
        
    # Data Rows
    for i_r, row_data in enumerate(data_rows):
        row = tbl.rows[i_r + 1]
        make_table_row_cant_split(row)
        bg_hex = "F8F9FA" if (i_r % 2 == 1) else "FFFFFF"
        for j, c_text in enumerate(row_data):
            cell = row.cells[j]
            cell.width = col_widths[j]
            set_cell_background(cell, bg_hex)
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            set_cell_borders(cell, top=('single', '4', '0', 'E8EAED'),
                                   bottom=('single', '4', '0', 'E8EAED'),
                                   left=('single', '4', '0', 'E8EAED'),
                                   right=('single', '4', '0', 'E8EAED'))
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            
            c_text_clean = c_text.replace('<br>', '\n').replace('<br/>', '\n')
            lines_in_cell = c_text_clean.split('\n')
            for l_idx, cell_line in enumerate(lines_in_cell):
                if l_idx > 0:
                    p = cell.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.1
                format_runs(p, cell_line, default_size=8.0)
                
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(2)
    p_after.paragraph_format.space_after = Pt(6)

def render_figure_box(doc, figure_lines):
    raw_text = "\n".join(figure_lines)
    m_title = re.search(r'<strong>(FIGURE [^<]+)</strong>', raw_text) or re.search(r'<div class="figure-title">([^<]+)</div>', raw_text)
    title = m_title.group(1) if m_title else "Executive Architecture & Flow Diagram"
    
    clean_lines = []
    if "FIGURE A.2" in title:
        clean_lines = [
            "**Target Market Sizing (2028)**: Total Addressable Market (TAM) = $21.0B | Serviceable Addressable Market (SAM) = $9.5B | Serviceable Obtainable Market (SOM) = $3.2B",
            "**Workload ARR Progression**: 2024: $275M ➔ 2026: $755M ➔ 2028 Target: $1,960M ($1.96B, doubling vertical share to 20%)",
            "**Workload Distribution (2028 Target)**: AI/ML & Gemini ($750M, 38%) · Data & OSDU ($520M, 27%) · HPC Seismic ($280M, 14%) · Sovereign & Security ($260M, 13%) · CCUS ($150M, 8%)"
        ]
    elif "FIGURE B.3" in title:
        clean_lines = [
            "**1. Ingestion**: OSDU Platform · SCADA Historians · LAS/DLIS Well Logs · Seismic 3D Volumes · P&ID Schematics",
            "**2. Gemini Reasoning**: 2M Token Context RAG · Subsurface & Timeseries FMs · Vertex Agent Builder · Anomaly Pattern Recognition",
            "**3. Governance Gate**: Human-in-the-Loop Sign-off · Geoscientist/HSE Gate · Deterministic Kill Switch · Complete Audit Trail Logging",
            "**4. Action Output**: Automated SAP Work Orders · SCADA Control Setpoints · Auto Petrel Surface Write · Emergency Rig Console Alerts"
        ]
    elif "FIGURE C.3" in title:
        clean_lines = [
            "**Dammam me-central2 (Class C Sovereign Region)**: CST / NCA Class C Certified · CNTXT External Key Management (EKM) · In-Country KSA Boundary",
            "**GDC Air-Gapped (OT/ICS SOC)**: Fully Disconnected Local Hardware · Mandiant ICS Threat Sensor · Local SCADA Zero-Trust SOC",
            "**HUMAIN AI Platform Engine (Saudi GTM)**: Saudi Aramco & Ministry AI · In-Kingdom Value Addition (IKVA) · SABIC & Ma'aden Enablement"
        ]
    elif "FIGURE D.1" in title:
        clean_lines = [
            "**Energy Operator Commitments (EQT, Pertamina, TotalEnergies)**: Firm Power PPAs (Gas/Geothermal/SMR) · Power-Adjacent Land & Permits · SCADA & Drilling Telemetry",
            "**Google & Alphabet Deliverables**: Preferential Anchor GCP Credits (20–40% below list) · DeepMind Science Lab Access · Tapestry Grid AI & Priority TPU/GPU SLAs",
            "**Bilateral Strategic Value**: Resolves data center power constraints while securing major GCP commitments with no competitor replication"
        ]
    else:
        text_only = re.sub(r'<[^>]+>', ' ', raw_text)
        items = [t.strip() for t in text_only.split('•') if t.strip()]
        clean_lines = items[:5] if items else ["Executive Architecture Diagram"]
        
    create_callout_box(doc, title, clean_lines, fill_hex="F8F9FA", border_color="1A73E8")

def build_word_doc(md_file_path, output_docx_path):
    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
        
    doc = docx.Document()
    
    # Page Setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        
        # Header & Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("Google Cloud · Global Oil & Gas Industry Strategy · Confidential")
        hrun.font.name = 'Arial'
        hrun.font.size = Pt(8)
        hrun.font.color.rgb = RGBColor(128, 134, 139)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        frun = fp.add_run("Strategic Executive Briefing · Page ")
        frun.font.name = 'Arial'
        frun.font.size = Pt(8)
        frun.font.color.rgb = RGBColor(128, 134, 139)
        
        # Add Page Number field XML to footer
        fldSimple = parse_xml(r'<w:fldSimple %s w:instr="PAGE"/>' % nsdecls('w'))
        fp._p.append(fldSimple)

    # Styles Setup
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(32, 33, 36) # #202124
    
    # Process lines
    lines = md_text.split('\n')
    i = 0
    total_lines = len(lines)
    
    table_rows = []
    
    while i < total_lines:
        line = lines[i].strip()
        
        # Handle Table parsing
        if line.startswith('|') and line.endswith('|'):
            table_rows.append(line)
            i += 1
            continue
        elif table_rows:
            render_markdown_table(doc, table_rows)
            table_rows = []
            
        if not line:
            i += 1
            continue
            
        if line == '---' or line == '***' or line == '___':
            p_sep = doc.add_paragraph()
            p_sep.paragraph_format.space_before = Pt(4)
            p_sep.paragraph_format.space_after = Pt(4)
            i += 1
            continue
            
        if '<div class="visual-card-container">' in line or '<div class="figure-card">' in line:
            figure_lines = []
            while i < total_lines and '</div>' not in lines[i]:
                figure_lines.append(lines[i])
                i += 1
            if i < total_lines:
                figure_lines.append(lines[i])
            i += 1
            render_figure_box(doc, figure_lines)
            continue
            
        if line.startswith('<div') or line.startswith('</div') or line.startswith('<span') or line.startswith('</span'):
            i += 1
            continue

        # Header 1
        if line.startswith('# '):
            h_text = line[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(26, 115, 232) # #1A73E8
            
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="4" w:color="1A73E8"/></w:pBdr>')
            p._p.get_or_add_pPr().append(pBdr)
            i += 1
            continue

        # Header 2
        if line.startswith('## '):
            h_text = line[3:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            run.font.name = 'Arial'
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(23, 78, 166) # #174EA6
            
            pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="3" w:color="DADCE0"/></w:pBdr>')
            p._p.get_or_add_pPr().append(pBdr)
            i += 1
            continue

        # Header 3
        if line.startswith('### '):
            h_text = line[4:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(32, 33, 36) # #202124
            i += 1
            continue

        # Header 4
        if line.startswith('#### '):
            h_text = line[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(h_text)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(60, 64, 67) # #3C4043
            i += 1
            continue

        # Bullet List Items
        if line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            format_runs(p, line[2:].strip())
            i += 1
            continue

        # Numbered List Items
        m_num = re.match(r'^(\d+)\.\s+(.*)$', line)
        if m_num:
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            format_runs(p, m_num.group(2).strip())
            i += 1
            continue

        # Standard Paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        format_runs(p, line)
        i += 1
        
    if table_rows:
        render_markdown_table(doc, table_rows)
        
    doc.save(output_docx_path)
    print(f"Successfully generated: {output_docx_path}")

if __name__ == '__main__':
    md_path = r'c:\Users\Husse\Documents\Google\google_cloud_og_industry_strategy.md'
    out_docx = r'c:\Users\Husse\Documents\Google\google_cloud_og_industry_strategy.docx'
    build_word_doc(md_path, out_docx)
